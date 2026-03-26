#!/usr/bin/env python3
"""
文档转换脚本 - 将 docx/doc/m文件转换为标准 Markdown 格式

用于将外部文档转换为幻协生存指南项目的标准文章格式。
"""

import sys
import os
import re
import argparse
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from PIL import Image
except ImportError:
    print("错误: 请先安装所需依赖")
    print("  uv pip install python-docx Pillow")
    sys.exit(1)


def extract_images_from_docx(doc, output_dir: str, article_title: str) -> dict:
    """从 docx 中提取图片并保存"""
    image_map = {}
    output_path = Path(output_dir) / article_title
    output_path.mkdir(parents=True, exist_ok=True)

    # 获取 docx 中的所有关系（包含图片）
    rels = doc.part.rels

    for rel_id, rel in rels.items():
        if "image" in rel.target_ref:
            image_data = rel.target_part.blob
            # 获取图片扩展名
            content_type = rel.target_part.content_type
            ext_map = {
                "image/png": ".png",
                "image/jpeg": ".jpg",
                "image/gif": ".gif",
                "image/webp": ".webp",
            }
            ext = ext_map.get(content_type, ".png")

            # 生成图片文件名
            img_name = f"image{rel_id}{ext}"
            img_path = output_path / img_name

            # 保存图片
            with open(img_path, "wb") as f:
                f.write(image_data)

            image_map[rel_id] = img_name

    return image_map


def extract_paragraphs(doc) -> list:
    """提取文档中的所有段落"""
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append({
                "text": text,
                "style": para.style.name if para.style else "Normal",
                "alignment": para.alignment,
            })
    return paragraphs


def extract_tables(doc) -> list:
    """提取文档中的所有表格"""
    tables = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        if table_data:
            tables.append(table_data)
    return tables


def convert_paragraph_to_markdown(para: dict) -> str:
    """将段落转换为 Markdown 格式"""
    text = para["text"]
    style = para["style"]

    # 根据样式判断标题级别
    if "Heading 1" in style or style == "标题 1":
        return f"# {text}"
    elif "Heading 2" in style or style == "标题 2":
        return f"## {text}"
    elif "Heading 3" in style or style == "标题 3":
        return f"### {text}"
    elif "Heading 4" in style or style == "标题 4":
        return f"#### {text}"
    else:
        return text


def convert_table_to_markdown(table_data: list) -> str:
    """将表格转换为 Markdown 格式"""
    if not table_data:
        return ""

    md_table = []
    for i, row in enumerate(table_data):
        # 转义管道符
        row = [cell.replace("|", "\\|") for cell in row]
        row_str = "| " + " | ".join(row) + " |"
        md_table.append(row_str)

        # 表头后添加分隔行
        if i == 0:
            md_table.append("|" + "|".join([" --- " for _ in row]) + "|")

    return "\n".join(md_table)


def convert_docx(
    input_file: str,
    output_file: str,
    title: str,
    author: str,
    school: str,
    date: str = None,
    categories: list = None,
    tags: list = None,
    image_dir: str = "draft/photos"
) -> str:
    """将 docx 文件转换为标准 Markdown 格式"""

    if categories is None:
        categories = ["幻协生存指南"]
    if tags is None:
        tags = []
    if date is None:
        date = datetime.now().strftime("%Y-%m-%d")

    # 打开文档
    doc = Document(input_file)

    # 提取内容
    paragraphs = extract_paragraphs(doc)
    tables = extract_tables(doc)

    # 提取图片
    image_map = {}
    if image_dir:
        image_map = extract_images_from_docx(doc, image_dir, title)

    # 构建 Markdown 内容
    content_lines = []

    # Front Matter
    content_lines.append("---")
    content_lines.append(f"title: {title}")
    content_lines.append(f"author: {author}")
    content_lines.append(f"date: {date}")
    content_lines.append("categories:")
    for cat in categories:
        content_lines.append(f"  - {cat}")
    if tags:
        content_lines.append("tags:")
        for tag in tags:
            content_lines.append(f"  - {tag}")
    content_lines.append("---")

    # HTML 包装结构
    content_lines.append('<section class="article">')
    content_lines.append('<header class="article-head">')
    content_lines.append(f'<p class="article-author">文 / {author} ({school})</p>')
    content_lines.append(f'<h1 class="article-title">{title}</h1>')
    content_lines.append("</header>")
    content_lines.append("")
    content_lines.append('<div class="article-body">')
    content_lines.append("")

    # 添加段落内容
    for para in paragraphs:
        md = convert_paragraph_to_markdown(para)
        content_lines.append(md)
        content_lines.append("")

    # 添加表格
    for table in tables:
        md_table = convert_table_to_markdown(table)
        content_lines.append(md_table)
        content_lines.append("")

    # 结束 HTML 标签
    content_lines.append("</div>")
    content_lines.append("</section>")

    # 写入文件
    output_path = Path(output_file)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n".join(content_lines))

    return output_file


def main():
    parser = argparse.ArgumentParser(
        description="将 docx 文档转换为标准 Markdown 格式"
    )
    parser.add_argument("input", help="输入的 docx 文件路径")
    parser.add_argument("-o", "--output", help="输出的 Markdown 文件路径")
    parser.add_argument("-t", "--title", required=True, help="文章标题")
    parser.add_argument("-a", "--author", required=True, help="作者名称")
    parser.add_argument("-s", "--school", required=True, help="作者学校/单位")
    parser.add_argument("-d", "--date", help="发布日期 (YYYY-MM-DD)")
    parser.add_argument("-c", "--categories", nargs="+", help="分类列表")
    parser.add_argument("--tags", nargs="+", help="标签列表")
    parser.add_argument("--image-dir", default="draft/photos", help="图片保存目录")

    args = parser.parse_args()

    # 设置输出文件路径
    if args.output:
        output_file = args.output
    else:
        output_file = f"draft/posts/{args.title}.md"

    # 转换文档
    try:
        result = convert_docx(
            input_file=args.input,
            output_file=output_file,
            title=args.title,
            author=args.author,
            school=args.school,
            date=args.date,
            categories=args.categories,
            tags=args.tags,
            image_dir=args.image_dir,
        )
        print(f"转换成功！文章已保存到: {result}")
        print(f"图片已保存到: {args.image_dir}/{args.title}/")
    except Exception as e:
        print(f"转换失败: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()
